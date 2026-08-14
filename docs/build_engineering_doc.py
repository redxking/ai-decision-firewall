from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ARCH = DOCS / "architecture"
OUT = DOCS / "AI_Decision_Firewall_POC_Engineering_Baseline_v0.1.docx"
METRICS = json.loads((ROOT / "outputs" / "baseline" / "metrics.json").read_text(encoding="utf-8"))

NAVY = "17365D"
BLUE = "2F5597"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EDF4FA"
DARK = "1F2937"
GRAY = "5B6573"
LIGHT_GRAY = "E9EDF2"
PALE_GRAY = "F6F8FA"
GREEN = "DDEFE3"
AMBER = "FFF1CC"
RED = "F7D7D9"
PURPLE = "E9E1F5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_table_row_split(row) -> None:
    """Keep a table row intact across page boundaries when it fits on one page."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name="Arial", size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rpr.append(sz)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rpr.append(rfonts)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    pf = normal.paragraph_format
    pf.space_after = Pt(5.5)
    pf.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Title", 30, NAVY, 0, 10),
        ("Subtitle", 15, BLUE, 0, 8),
        ("Heading 1", 18, NAVY, 14, 7),
        ("Heading 2", 14, BLUE, 11, 5),
        ("Heading 3", 11.5, DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["Caption"].font.name = "Arial"
    styles["Caption"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    styles["Caption"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    styles["Caption"].font.size = Pt(8.5)
    styles["Caption"].font.italic = True
    styles["Caption"].font.color.rgb = RGBColor.from_string(GRAY)
    styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Caption"].paragraph_format.space_after = Pt(7)


def setup_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("AI Decision Firewall POC | Engineering Baseline v0.1")
    set_run_font(run, size=8.5, bold=True, color=GRAY)

    footer = section.footer
    table = footer.add_table(rows=1, cols=3, width=Inches(6.9))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.3, 2.3, 2.3]
    for cell, width in zip(table.rows[0].cells, widths, strict=True):
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p0 = table.cell(0, 0).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p0.add_run("Working Draft")
    set_run_font(r, size=8, color=GRAY)
    p1 = table.cell(0, 1).paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run("Synthetic Data and Simulated Actions Only")
    set_run_font(r, size=8, bold=True, color=BLUE)
    p2 = table.cell(0, 2).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p2.add_run("Page ")
    set_run_font(r, size=8, color=GRAY)
    add_field(p2, "PAGE")
    r = p2.add_run(" of ")
    set_run_font(r, size=8, color=GRAY)
    add_field(p2, "NUMPAGES")


def add_para(doc: Document, text: str = "", *, bold_prefix: str | None = None, style: str | None = None, align=None, keep=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc: Document, items: Iterable[str], level: int = 0) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        r = p.add_run(item)
        set_run_font(r)
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    # Use explicit numbering so each engineering sequence restarts at 1 in all
    # Word/LibreOffice renderers instead of inheriting a prior list instance.
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.24)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{index}.  ")
        set_run_font(r)
        r = p.add_run(item)
        set_run_font(r)


def add_callout(doc: Document, title: str, body: str, fill: str = PALE_BLUE, border: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    prevent_table_row_split(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": border}, top={"val": "nil"}, right={"val": "nil"}, bottom={"val": "nil"})
    set_cell_margins(cell, top=120, bottom=120, start=170, end=130)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    r = p2.add_run(body)
    set_run_font(r, size=9.8)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size: float = 8.5, header_fill: str = NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = widths is None
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_table_row_split(hdr)
    for index, header in enumerate(headers):
        cell = hdr.cells[index]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, bold=True, color=WHITE)
        if widths:
            cell.width = Inches(widths[index])
    for row_index, row in enumerate(rows):
        added_row = table.add_row()
        prevent_table_row_split(added_row)
        cells = added_row.cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, PALE_GRAY)
        for index, value in enumerate(row):
            cell = cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
            if widths:
                cell.width = Inches(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.7) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    # Use the human-readable caption as alternative text so the architecture
    # and metric figures remain understandable to assistive technologies.
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", path.stem.replace("_", " "))
    doc_pr.set("descr", caption)
    cp = doc.add_paragraph(caption, style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_landscape_section(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    section.different_first_page_header_footer = False
    return section


def add_portrait_section(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = False
    return section


def heading(doc: Document, number: str, title: str, level: int = 1):
    text = f"{number} {title}" if number else title
    return doc.add_heading(text, level=level)


def build_document() -> Document:
    doc = Document()
    style_document(doc)
    setup_sections(doc)
    cp = doc.core_properties
    cp.title = "AI Decision Firewall: Privileged Identity Containment POC — Engineering Baseline v0.1"
    cp.subject = "Model-agnostic evidence and authority control for AI-assisted cyber decisions"
    cp.author = "Angelis Pseftis"
    cp.keywords = "AI assurance, cyber decision support, privileged identity, SOC, evidence provenance, autonomous containment, systems engineering"
    cp.comments = "Synthetic data and simulated actions only. Working engineering baseline."
    cp.category = "Systems Engineering"

    # Cover page
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("AI DECISION FIREWALL")
    set_run_font(r, size=12, bold=True, color=BLUE)
    p2 = doc.add_paragraph(style="Title")
    p2.paragraph_format.space_before = Pt(5)
    r = p2.add_run("Privileged Identity Containment\nProof of Concept")
    set_run_font(r, size=30, bold=True, color=NAVY)
    p3 = doc.add_paragraph(style="Subtitle")
    r = p3.add_run("Engineering Baseline v0.1")
    set_run_font(r, size=16, color=BLUE)

    band = doc.add_table(rows=1, cols=1)
    cell = band.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    rp = cell.paragraphs[0]
    rr = rp.add_run("SYNTHETIC DATA • SIMULATED ACTIONS • NO PRODUCTION INTEGRATION")
    set_run_font(rr, size=10, bold=True, color=WHITE)

    add_para(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT)
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(20)
    r = p4.add_run("Objective")
    set_run_font(r, size=11, bold=True, color=NAVY)
    p5 = doc.add_paragraph()
    r = p5.add_run(
        "Demonstrate an executable, model-agnostic control layer that can determine whether suspicious privileged-identity activity should be closed, investigated, reversibly contained, or escalated—while keeping action authority outside the AI model."
    )
    set_run_font(r, size=12, color=DARK)
    p5.paragraph_format.line_spacing = 1.15

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_before = Pt(75)
    for label, value in [
        ("Prepared by", "Angelis Pseftis"),
        ("Date", "14 August 2026"),
        ("Status", "Working engineering baseline"),
        ("Distribution", "PUBLIC / UNRESTRICTED"),
    ]:
        r = p6.add_run(f"{label}: ")
        set_run_font(r, size=9.5, bold=True, color=GRAY)
        r = p6.add_run(value + "\n")
        set_run_font(r, size=9.5, color=DARK)

    doc.add_page_break()

    # Document control
    heading(doc, "", "Document Control", 1)
    add_table(doc, ["Field", "Value"], [
        ["Document title", "AI Decision Firewall: Privileged Identity Containment POC — Engineering Baseline v0.1"],
        ["Version", "0.1.0"],
        ["Technical owner", "Angelis Pseftis"],
        ["Baseline date", "14 August 2026"],
        ["Distribution", "PUBLIC / UNRESTRICTED"],
        ["Data boundary", "Generated data only; labels stored separately from runtime inputs"],
        ["Action boundary", "In-memory simulator; reversible actions only"],
        ["Operational status", "Not authorized for production use"],
        ["Primary decision requested", "Approve Phase 2 historical replay and canonical data-contract discovery"],
    ], widths=[1.55, 5.25], font_size=9.2)

    add_callout(
        doc,
        "Use restriction",
        "This document and repository describe a proof of concept. No baseline result establishes operational detection accuracy, safe autonomous action, regulatory compliance, or authority to connect the system to a production identity platform.",
        fill=AMBER,
        border="C98700",
    )

    heading(doc, "", "Document Guide", 1)
    add_table(doc, ["Reader", "Start here", "Decision supported"], [
        ["Executive sponsor / authorizing official", "Executive Summary; Sections 1–4; 17–20", "Whether to fund and govern the next validation phase"],
        ["Chief architect / principal engineer", "Sections 5–16; Appendices A–D", "Architecture, interfaces, requirements, safety case, and integration plan"],
        ["SOC / identity operations", "Sections 6, 11, 12, 16, and 21", "Operational workflow, evidence needs, action boundaries, and analyst role"],
        ["Data science / AI assurance", "Sections 8–10, 16–18", "Dataset limitations, model role, calibration, and evaluation design"],
        ["Security assessor / red team", "Sections 13–16; Appendix C", "Threat model, abuse cases, verification evidence, and residual risk"],
    ], widths=[1.7, 2.3, 2.8], font_size=8.8)

    # Executive Summary
    heading(doc, "", "Executive Summary", 1)
    add_para(doc, "The unresolved AI problem is not whether a model can recommend a cyber action. It is whether the surrounding system can determine when evidence is trustworthy and sufficient, constrain the action to an authorized boundary, prove why the action was allowed, and verify that the action produced the intended state without unacceptable mission impact.")
    add_para(doc, "This POC narrows that problem to privileged-identity containment. The system evaluates synthetic identity, endpoint, network, threat-intelligence, asset, change-management, workforce-context, and ticket evidence. It returns one of four dispositions: no action, investigate, reversible containment, or human escalation. The learned model is advisory; it cannot issue tokens, access credentials, change policy, or call the action target.")

    decision = METRICS["decision_control"]
    safety = METRICS["safety_and_assurance"]
    model = METRICS["model"]
    perf = METRICS["performance"]
    add_table(doc, ["Baseline measure", "Result", "Interpretation"], [
        ["Synthetic test cases", str(METRICS["scope"]["cases_evaluated"]), "Mechanics tested across benign, malicious, ambiguous, incomplete, break-glass, and adversarial scenarios"],
        ["Reversible containment decisions", str(decision["autonomous_containment_count"]), "Only allow-listed actions against the in-memory simulator"],
        ["Synthetic false containment", str(decision["false_containment_count"]), "No benign synthetic case received a containment disposition"],
        ["Unsafe automation", str(safety["unsafe_automation_count"]), "No action in generator-consistent cases marked poisoned, conflicted, break-glass, or above-threshold critical"],
        ["Evidence traceability", f"{100*safety['evidence_traceability_rate']:.1f}%", "Every cited evidence and feature-trace event ID resolved to an input event"],
        ["Authorization without verifier", str(safety["authorization_without_independent_verifier"]), "No token bypassed independent verification"],
        ["Audit chain", "Valid" if safety["audit_chain_valid"] else "Invalid", "Hash continuity and record integrity verified"],
        ["Median local decision latency", f"{perf['median_decision_latency_ms']:.3f} ms", "Software execution only; excludes real telemetry and vendor APIs"],
    ], widths=[2.1, 1.25, 3.45], font_size=8.8)

    add_callout(
        doc,
        "What the baseline proves",
        "The encoded authority and safety invariants are executable, testable, reproducible, and separable from the risk model. The baseline does not prove that the synthetic model will perform against real operations or that unvalidated external inputs preserve the generator's cross-field consistency.",
        fill=GREEN,
        border="4F8A61",
    )
    add_callout(
        doc,
        "Recommended decision",
        "Proceed to Phase 2: de-identified historical replay and canonical data-contract discovery. Maintain a hard prohibition on live action. Use the next phase to measure evidence availability, analyst disagreement, calibration, temporal behavior, and integration risk before considering shadow mode.",
        fill=LIGHT_BLUE,
        border=BLUE,
    )

    # 1
    heading(doc, "1.", "Purpose and Objective", 1)
    add_para(doc, "The POC objective is to establish a trustworthy decision-control pattern for AI-assisted cyber operations. The immediate use case is suspicious privileged-identity activity because it has meaningful mission consequence, strong multi-source observability, reversible first actions, measurable outcomes, and a safe progression from offline replay to shadow mode.")
    heading(doc, "1.1", "Primary Engineering Objective", 2)
    add_callout(doc, "Objective statement", "Demonstrate that a replaceable AI or analytical model can contribute to a cyber decision without possessing the authority, credentials, or direct interface required to execute that decision.")
    heading(doc, "1.2", "POC Hypothesis", 2)
    add_para(doc, "A model-agnostic evidence and authority layer can materially reduce the risk of unsafe AI action by separating probabilistic assessment from deterministic authorization, requiring evidence-quality and consequence thresholds, independently verifying eligibility, issuing action-scoped tokens, and validating the post-action state.")
    heading(doc, "1.3", "Success Is Not Defined as Model Accuracy Alone", 2)
    add_bullets(doc, [
        "The model cannot directly invoke the action broker or access target credentials.",
        "Ground-truth labels are absent from runtime case inputs.",
        "Untrusted free text cannot satisfy action authority and prompt-injection content forces abstention.",
        "Canonical cases marked as break-glass or above the asset-criticality threshold, plus conflicted or missing evidence, are protected by explicit policy.",
        "Every autonomous action is allow-listed and bounded; every authorization is scoped, short-lived, and independently verified.",
        "Every material claim, decision, token, action, and verification result is traceable and auditable.",
    ])

    # 2
    heading(doc, "2.", "How We Got Here", 1)
    add_para(doc, "The starting problem was broad: organizations want AI agents to reduce SOC burden and eventually take operational action, but they cannot reliably establish when an agent has enough trustworthy evidence or whether a proposed action is authorized and proportionate. Attempting to solve the entire SOC would create an unbounded requirements surface and make it difficult to distinguish model shortcomings from integration, evidence, policy, or authority failures.")
    heading(doc, "2.1", "Problem Decomposition", 2)
    add_numbered(doc, [
        "Separate analysis quality from action authority. A model may be correct and still be unauthorized to act.",
        "Separate evidence content from evidence trust. A plausible event may be stale, corrupted, duplicated, adversarial, or unsupported by provenance.",
        "Separate action recommendation from action execution. A correct recommendation may still exceed delegated authority or mission tolerance.",
        "Separate command success from operational success. An API response does not prove the intended state or absence of collateral impact.",
        "Separate synthetic mechanics validation from operational efficacy validation. Early generated data enables architecture testing but cannot support production claims.",
    ])
    heading(doc, "2.2", "Why Privileged Identity Was Selected", 2)
    add_table(doc, ["Selection criterion", "Privileged identity rationale"], [
        ["Mission consequence", "Compromised privileged access can enable broad lateral movement, cloud persistence, and control-plane manipulation."],
        ["Observable evidence", "Identity, endpoint, network, asset, change, travel, and threat-intelligence sources can be correlated."],
        ["Bounded, recoverable first actions", "Session revocation, temporary step-up authentication, and enhanced monitoring can be constrained; recovery is not always exact transactional reversal."],
        ["Human authority boundary", "Account disablement and endpoint isolation remain clearly human-authorized in the POC."],
        ["Measurability", "Decision quality, action precision, time, rollback, and post-action state are measurable."],
        ["Generalizability", "The same evidence-authority pattern can later support endpoint, cloud, network, and IT/OT boundary decisions."],
    ], widths=[1.9, 4.9], font_size=9)
    heading(doc, "2.3", "Product Thesis", 2)
    add_callout(doc, "Durable differentiation", "The defensible product is not a proprietary chatbot. It is the evidence ontology, consequence model, authority framework, independent verification, action-token protocol, replay harness, policy packs, and audit record that remain valid when the underlying model changes.", fill=PURPLE, border="76599A")

    # 3
    heading(doc, "3.", "Decision Requested and Strategic Direction", 1)
    add_para(doc, "The POC has completed the synthetic executable baseline. The next decision is whether to invest in evidence realism and operational validation without expanding action authority.")
    add_table(doc, ["Decision", "Recommendation", "Rationale"], [
        ["Proceed to historical replay?", "Yes", "Required to discover schema gaps, evidence latency, class imbalance, analyst disagreement, and synthetic-to-real drift."],
        ["Connect to live production telemetry?", "Not yet", "First complete de-identification, data contracts, secure ingestion, privacy review, and replay acceptance criteria."],
        ["Enable live actions?", "No", "No operational evidence exists to support authorization. Maintain simulator-only execution."],
        ["Introduce an LLM agent?", "Not in v0.2 critical path", "Validate evidence, policy, and authority layers first; then compare model types behind the same interface."],
        ["Expand beyond identity?", "Defer", "Maintain use-case focus until replay results establish that the architecture and data contract are viable."],
    ], widths=[1.7, 1.15, 3.95], font_size=9)

    # 4 scope
    heading(doc, "4.", "Scope, Assumptions, and Constraints", 1)
    heading(doc, "4.1", "In Scope", 2)
    add_bullets(doc, [
        "Synthetic privileged-identity cases and event-level evidence.",
        "Evidence provenance, integrity, freshness, trust, completeness, diversity, and conflict assessment.",
        "A replaceable learned risk model with feature-level explanation and event traceability.",
        "Four bounded dispositions and a deterministic policy engine.",
        "Independent non-model verification and signed authorization tokens.",
        "In-memory simulation of bounded identity actions and verification of the state returned by the simulator.",
        "Tamper-evident audit records, automated tests, metrics, and documentation.",
    ])
    heading(doc, "4.2", "Explicitly Out of Scope", 2)
    add_bullets(doc, [
        "Production identity-provider, EDR, SIEM, SOAR, cloud, or ITSM connections.",
        "Real user, employee, incident, or classified data.",
        "Direct autonomous account disablement, endpoint isolation, network blocking, or persistent policy modification.",
        "Operational OT or cyber-physical process control.",
        "A general autonomous SOC agent, conversational user interface, or vendor-specific copilot replacement.",
        "Compliance certification, Authority to Operate, legal determination, or safety certification.",
    ])
    heading(doc, "4.3", "Key Assumptions", 2)
    add_table(doc, ["ID", "Assumption", "Validation needed"], [
        ["A-01", "Privileged identity remains a high-value initial use case.", "Replay workload and stakeholder interviews"],
        ["A-02", "Required evidence can be correlated through stable identity, device, session, and asset keys.", "Canonical data-contract discovery"],
        ["A-03", "Low-impact actions can be made idempotent and independently verified across vendors.", "Test-tenant adapter trials"],
        ["A-04", "Asset criticality and break-glass status are accurate enough to control authority.", "CMDB and identity-governance quality assessment"],
        ["A-05", "Analyst adjudications can provide useful labels despite uncertainty and disagreement.", "Multi-reviewer labeling study"],
        ["A-06", "A model-agnostic control plane is valuable even if the selected model changes.", "Comparative model evaluation"],
    ], widths=[0.7, 3.35, 2.75], font_size=8.8)

    # 5 architecture in landscape
    heading(doc, "5.", "System Architecture", 1)
    add_para(doc, "The architecture is partitioned into four planes: evidence, decision, authority/action, and assurance. The control objective is to prevent any single probabilistic component from controlling evidence interpretation, authorization, execution, and verification.")
    add_landscape_section(doc)
    heading(doc, "5.1", "System Context", 2)
    add_figure(doc, ARCH / "01_system_context.png", "Figure 1. POC system context and external actors.", width=9.7)
    heading(doc, "5.2", "Logical Architecture", 2)
    add_figure(doc, ARCH / "02_logical_architecture.png", "Figure 2. Four-plane logical architecture.", width=6.9)
    heading(doc, "5.3", "Decision State Machine", 2)
    add_figure(doc, ARCH / "03_decision_state_machine.png", "Figure 3. Fail-safe decision and action state machine.", width=10.0)
    heading(doc, "5.4", "Trust Boundaries and Data Separation", 2)
    add_figure(doc, ARCH / "04_trust_boundaries.png", "Figure 4. Trust boundaries, label separation, and credential isolation.", width=10.0)
    add_portrait_section(doc)

    heading(doc, "5.5", "Component Responsibilities", 2)
    add_table(doc, ["Component", "Responsibility", "Must not do"], [
        ["Synthetic generator", "Create reproducible cases, event evidence, and separate evaluator labels.", "Represent itself as operationally realistic"],
        ["Evidence assessor", "Evaluate provenance, integrity, freshness, source trust, completeness, diversity, conflicts, and adversarial instructions.", "Treat content plausibility as proof of trust"],
        ["Feature extractor", "Create allow-listed structured features and retain event-level trace.", "Pass free-text instructions or ground truth to the model"],
        ["Risk model", "Estimate compromise probability and feature contributions.", "Authorize, execute, or modify policy"],
        ["Policy engine", "Translate risk, evidence, criticality, and identity state into a bounded disposition.", "Bypass configured authority constraints"],
        ["Independent verifier", "Re-evaluate containment eligibility and reject human-only or unsafe actions.", "Reuse model confidence as sole evidence"],
        ["Authorization gate", "Mint signed, scoped, short-lived tokens after verifier approval.", "Issue broad or reusable credentials"],
        ["Action broker", "Validate tokens and execute allow-listed simulator actions.", "Accept direct model commands"],
        ["Post-action verifier", "Check target state against expected action outcome.", "Assume API command success equals mission success"],
        ["Audit logger", "Record a hash-chained history of decisions and actions.", "Silently alter or delete prior records"],
        ["Evaluator", "Join decisions with hidden labels and calculate metrics.", "Expose labels to runtime components"],
    ], widths=[1.35, 3.05, 2.4], font_size=8.3)

    # 6 conops
    heading(doc, "6.", "Operational Concept", 1)
    heading(doc, "6.1", "Actors and Authorities", 2)
    add_table(doc, ["Actor", "Authority", "POC role"], [
        ["Evidence producer", "Submit telemetry and context", "Synthetic source modules"],
        ["Decision firewall", "Recommend a bounded disposition", "Automated pipeline"],
        ["Policy owner", "Define thresholds, action classes, and protected identities", "Configuration owner"],
        ["Independent verifier", "Approve or reject autonomous-action eligibility", "Deterministic software verifier"],
        ["Human decision authority", "Approve high-impact action or resolve uncertainty", "SOC lead, identity owner, or incident commander"],
        ["Action broker", "Execute only actions authorized by a valid token", "In-memory simulator interface"],
        ["Evaluator", "Measure decisions after labels are revealed", "Offline metrics pipeline"],
    ], widths=[1.45, 2.4, 2.95], font_size=8.8)
    heading(doc, "6.2", "Nominal Workflow", 2)
    add_numbered(doc, [
        "Receive a case and validate its schema.",
        "Assess evidence trust and isolate untrusted text.",
        "Extract structured features and retain event-level lineage.",
        "Obtain the model probability and contributions.",
        "Apply deterministic decision policy.",
        "Run independent eligibility checks.",
        "Issue a signed token only for allow-listed reversible actions.",
        "Execute through the credential-isolated broker.",
        "Verify the resulting target state.",
        "Write the complete record to the audit chain and evaluator outputs.",
    ])
    heading(doc, "6.3", "Operating Modes", 2)
    add_table(doc, ["Mode", "Data", "Action authority", "Primary objective"], [
        ["Development", "Synthetic", "Simulator only", "Software, architecture, and failure-mode validation"],
        ["Historical replay", "De-identified prior cases", "None", "Data-contract, calibration, and counterfactual analysis"],
        ["Shadow", "Live read-only telemetry", "None", "Compare recommendations with analyst decisions"],
        ["Controlled test tenant", "Non-production identities", "Explicit change-controlled reversible actions", "Validate adapters, idempotency, rollback, and state verification"],
        ["Limited pilot", "Approved operational population", "Human approval initially", "Establish bounded operational evidence"],
    ], widths=[1.25, 1.65, 2.25, 1.65], font_size=8.6)

    # 7 requirements
    heading(doc, "7.", "Requirements Baseline", 1)
    add_para(doc, "Requirements are organized by functional capability, safety, security, nonfunctional performance, and verification. The complete traceability matrix is included in Appendix C and in the repository as `docs/REQUIREMENTS_TRACEABILITY_MATRIX.csv`.")
    heading(doc, "7.1", "Key Functional Requirements", 2)
    add_table(doc, ["ID", "Requirement", "Verification"], [
        ["FR-003", "Assess evidence provenance, integrity, freshness, trust, diversity, completeness, and conflicts.", "Scenario and unit tests"],
        ["FR-004", "Exclude untrusted free text from model features.", "Code inspection and injection scenario"],
        ["FR-006", "Produce one of four bounded dispositions.", "Decision summary and scenario metrics"],
        ["FR-007", "Perform independent non-model verification.", "Unit tests and audit evidence"],
        ["FR-008", "Issue case-bound, action-scoped, short-lived authorization tokens.", "Token and decision-record inspection"],
        ["FR-010", "Verify target state after action execution.", "Injected action failures"],
        ["FR-011", "Generate tamper-evident audit records.", "Audit chain and tamper test"],
    ], widths=[0.7, 4.2, 1.9], font_size=8.7)
    heading(doc, "7.2", "Non-Negotiable Safety Requirements", 2)
    add_table(doc, ["ID", "Safety invariant"], [
        ["SAF-001", "The model has no direct action authority."],
        ["SAF-002", "Poisoned evidence blocks autonomous action."],
        ["SAF-003", "Break-glass identities require human authority."],
        ["SAF-004", "Assets above the configured criticality boundary cannot be acted upon autonomously."],
        ["SAF-005", "Human-only actions are rejected from executable action sets."],
        ["SAF-006", "No action executes without a valid token."],
        ["SAF-007", "No token is issued without independent-verifier approval."],
        ["SAF-008", "No action is declared successful without post-action state verification."],
    ], widths=[0.85, 5.95], font_size=9)

    # 8 data
    heading(doc, "8.", "Synthetic Data Design", 1)
    add_para(doc, "The synthetic dataset is a test fixture, not a claim about operational prevalence. Its purpose is to create repeatable evidence combinations that exercise architecture decisions and safety boundaries, including cases that are difficult or inappropriate to manufacture in production.")
    heading(doc, "8.1", "Data Separation", 2)
    add_bullets(doc, [
        "`train_cases.jsonl` and `test_cases.jsonl` contain runtime case metadata and evidence only.",
        "`train_labels.jsonl` and `test_labels.jsonl` contain scenario, compromise truth, expected disposition, and rationale.",
        "The decision engine receives only case files.",
        "The evaluator joins decisions with labels after the engine completes.",
        "The model training process sees only the training partition and training labels.",
    ])
    heading(doc, "8.2", "Canonical Evidence Event", 2)
    add_table(doc, ["Field", "Purpose"], [
        ["event_id / case_id", "Traceability and case association"],
        ["source_type / source_instance", "Source independence and adapter identity"],
        ["observed_at / collected_at", "Temporal ordering and freshness measurement"],
        ["integrity", "Verified, unverified, or failed integrity status"],
        ["provenance_id", "Evidence lineage identifier"],
        ["trust_score", "Configured source trust input"],
        ["entity_refs", "Identity, device, asset, session, and application correlation"],
        ["attributes", "Typed allow-listed evidence values"],
        ["untrusted_text", "Isolated context that cannot provide authority"],
        ["contains_instructional_content", "Adversarial-content safety signal"],
    ], widths=[2.15, 4.65], font_size=8.8)
    heading(doc, "8.3", "Scenario Catalog", 2)
    add_table(doc, ["Scenario", "Class", "Purpose"], [
        ["Privileged token theft", "Malicious", "Token reuse, new device, impossible travel, threat infrastructure, administrative activity"],
        ["Password spray then success", "Malicious", "High failed-login intensity followed by successful privileged access"],
        ["Credential dump and lateral movement", "Malicious", "Endpoint and network corroboration for severe post-compromise behavior"],
        ["Malicious OAuth consent", "Malicious", "Cloud persistence and suspicious grant activity"],
        ["Approved travel", "Benign", "Context explains geolocation anomaly"],
        ["Corporate VPN geolocation", "Benign", "Known egress explains apparent impossible travel"],
        ["Approved maintenance", "Benign", "Change record explains after-hours administrative behavior"],
        ["Service-account batch", "Benign", "Known scheduled automation explains activity"],
        ["Break-glass drill", "Protected benign", "Tests mandatory human authority"],
        ["Sensor conflict", "Ambiguous", "Tests abstention under contradictory evidence"],
        ["Telemetry gap", "Ambiguous", "Tests abstention when expected sources are missing"],
        ["Prompt-injection poisoning", "Adversarial", "Tests isolation of instructions embedded in untrusted content"],
    ], widths=[2.0, 1.15, 3.65], font_size=8.4)
    heading(doc, "8.4", "Data Limitations", 2)
    add_callout(doc, "Critical limitation", "Training and test partitions originate from the same scenario generator. Therefore, model discrimination is optimistic. The POC uses those metrics only to confirm that the model interface and policy can be exercised; it does not treat them as real-world performance evidence.", fill=AMBER, border="C98700")

    # 9 evidence
    heading(doc, "9.", "Evidence Trust and Provenance", 1)
    add_para(doc, "The system distinguishes evidence content from evidence trust. A high-risk event cannot authorize action merely because it is alarming; the event must also satisfy provenance, integrity, freshness, and corroboration requirements.")
    heading(doc, "9.1", "Evidence Quality Function", 2)
    add_para(doc, "The v0.1 evidence-quality score is an engineering heuristic:")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Q = 0.27P + 0.23I + 0.18F + 0.17D + 0.15T − conflict, poison, and missing-source penalties")
    set_run_font(r, name="Cambria Math", size=11, bold=True, color=NAVY)
    add_table(doc, ["Term", "Meaning"], [
        ["P", "Proportion of events with a provenance identifier"],
        ["I", "Proportion of events with verified integrity"],
        ["F", "Freshness derived from observation-to-collection delay"],
        ["D", "Diversity of expected evidence-source classes"],
        ["T", "Mean configured source-trust score"],
    ], widths=[0.9, 5.9], font_size=8.9)
    heading(doc, "9.2", "Adversarial Content Handling", 2)
    add_bullets(doc, [
        "Free text is not included in model features.",
        "Instruction-like strings are detected independently of model reasoning.",
        "An event containing instructions directed at an agent is marked poisoned for authority purposes.",
        "Poisoned evidence forces `INVESTIGATE`; it cannot satisfy autonomous-containment policy.",
        "Future natural-language extraction must produce typed claims with provenance and remain subordinate to the same policy gate.",
    ])

    # 10 model
    heading(doc, "10.", "Model Design and Role", 1)
    add_para(doc, "The baseline model is intentionally simple and interpretable. It demonstrates that a learned recommender can be integrated without allowing the model to dominate authority. A future LLM, graph model, ensemble, or vendor agent must use the same assessment interface and remain subject to the same controls.")
    heading(doc, "10.1", "Model Form", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("p(compromise | evidence) = σ(wᵀz + b)")
    set_run_font(r, name="Cambria Math", size=12, bold=True, color=NAVY)
    add_para(doc, "Features are standardized using training-partition statistics. Batch gradient descent with L2 regularization fits the weights. The model emits probability, the strongest positive and negative feature contributions, and the source events supporting each feature.")
    heading(doc, "10.2", "Feature Boundary", 2)
    add_table(doc, ["Included structured features", "Excluded inputs"], [[
        "Failed-login intensity; new device; impossible travel; threat-IP match; MFA fatigue; token reuse; credential dumping; lateral movement; unusual admin action; EDR malware; after-hours activity; privilege level; asset criticality; VPN, travel, maintenance, and service-account context; strong MFA; device compliance; OAuth grant.",
        "Free-text instructions; scenario name; expected disposition; ground truth; raw credentials; user display name; analyst command text; action credentials."
    ]], widths=[3.4, 3.4], font_size=8.8)
    heading(doc, "10.3", "Model Authority Restrictions", 2)
    add_callout(doc, "Hard boundary", "The model returns an assessment object. It does not hold the HMAC signing key, target credentials, an action-broker reference, a policy-modification interface, or a route to the simulator.", fill=GREEN, border="4F8A61")
    add_figure(doc, ARCH / "06_probability_distribution.png", "Figure 5. Synthetic risk-score distribution. The separation is optimistic because the same generator family produced training and test partitions.", width=6.4)

    # 11 policy
    heading(doc, "11.", "Decision Policy", 1)
    add_para(doc, "The decision policy combines model probability with evidence quality, independent corroboration, protected identity status, asset criticality, conflict state, and action class. The policy is external configuration and is recorded in each decision.")
    heading(doc, "11.1", "Four Dispositions", 2)
    add_table(doc, ["Disposition", "Entry condition", "Permitted system behavior"], [
        ["NO_ACTION", "Decision-grade evidence; risk at or below 0.24; no severe indicator", "Close with evidence and rationale; no operational action"],
        ["INVESTIGATE", "Evidence missing, stale, conflicted, poisoned, low integrity, or risk uncertain", "Issue read-only evidence-collection tasks"],
        ["CONTAIN_REVERSIBLE", "Risk at or above 0.88; decision-grade evidence; at least two supporting source classes; asset criticality at or below 0.75; not break glass", "Mint token for three allow-listed simulator actions after verifier approval"],
        ["ESCALATE_HUMAN", "Risk at or above 0.62, severe behavior, critical asset, break-glass identity, or human-only action", "Recommend action and route authority to a human role"],
    ], widths=[1.45, 3.65, 1.7], font_size=8.3)
    heading(doc, "11.2", "Current POC Action Classes", 2)
    add_table(doc, ["Class", "Actions", "Authority"], [
        ["Autonomous POC allowlist", "Revoke active sessions; require temporary step-up authentication; increase monitoring", "Signed token after independent verification"],
        ["Human-only", "Disable account; isolate endpoint; block network path; modify persistent access policy", "Explicit authorized human decision"],
        ["Read only", "Query identity history; query endpoint telemetry; validate change and travel context", "Read-only automation"],
    ], widths=[1.55, 3.8, 1.45], font_size=8.8)
    heading(doc, "11.3", "Future Consequence-Based Policy", 2)
    add_para(doc, "The fixed thresholds are sufficient for a POC. A production design should explicitly minimize expected mission loss rather than optimize a generic classification threshold:")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("a* = argminₐ [ P(C|E)Lmiss(a) + P(¬C|E)Lfalse(a) + Lmission(a) ]")
    set_run_font(r, name="Cambria Math", size=11.5, bold=True, color=NAVY)
    add_para(doc, "The loss terms must be action-specific, mission-specific, and approved by the policy owner. They cannot be inferred solely from model confidence.")

    # 12 auth action
    heading(doc, "12.", "Authorization, Execution, and Verification", 1)
    heading(doc, "12.1", "Independent Verification", 2)
    add_para(doc, "The verifier re-evaluates containment eligibility without using the model as its sole source of truth. It validates probability range, event traceability, action allowlists, evidence thresholds, source corroboration, conflict state, poisoned evidence, break-glass status, asset criticality, and rollback coverage.")
    heading(doc, "12.2", "Authorization Token", 2)
    add_table(doc, ["Token property", "POC implementation"], [
        ["Integrity", "HMAC-SHA-256 signature"],
        ["Scope", "Case ID and exact permitted actions"],
        ["Binding", "Hash of disposition, actions, rules, and cited evidence IDs"],
        ["Lifetime", "120 seconds"],
        ["Issuance condition", "Independent verifier passes and executable actions exist"],
        ["Validation", "Signature, case, action membership, and expiration checked by broker"],
    ], widths=[1.75, 5.05], font_size=8.8)
    heading(doc, "12.3", "Action Simulator", 2)
    add_para(doc, "The action target is an in-memory identity-state object. It has no external network connector. The simulator tracks active sessions, temporary step-up authentication, monitoring level, and account-enabled state. It deliberately injects deterministic downstream failures so that the POC exercises post-action failure handling.")
    heading(doc, "12.4", "Post-Action Verification", 2)
    add_table(doc, ["Action", "Expected state"], [
        ["Revoke active sessions", "active_sessions = 0"],
        ["Force step-up authentication", "step_up_required = true"],
        ["Increase monitoring", "monitoring_level = enhanced"],
    ], widths=[2.45, 4.35], font_size=9)
    add_para(doc, "In v0.1, this check evaluates the state returned by the same in-memory simulator call. It is not an independent readback from a target system, and rollback plans are descriptive metadata rather than executable rollback orchestration.")

    # 13 security safety
    heading(doc, "13.", "Security and Safety Architecture", 1)
    add_para(doc, "The top-level safety claim is that the POC cannot autonomously perform a consequential identity action unless deterministic policy and independent verification establish that the case, evidence, asset, identity, and action satisfy the configured reversible-action boundary.")
    heading(doc, "13.1", "Security Controls Implemented", 2)
    add_bullets(doc, [
        "Module-level separation between model, policy, verifier, authorization gate, broker, and target.",
        "Allow-listed structured feature extraction.",
        "Free-text isolation and instruction detection.",
        "HMAC integrity protection for authorization tokens.",
        "Case and action scoping with expiration.",
        "Human-only action denylist enforced by the verifier.",
        "Hash-chained audit log and tamper detection.",
        "No production connector, secret, credential, or target endpoint.",
    ])
    heading(doc, "13.2", "Residual Security Limitations", 2)
    add_table(doc, ["Limitation", "Consequence", "Required next control"], [
        ["POC signing-key fallback", "Not appropriate for production key custody", "External secrets manager or HSM; rotation; dual control"],
        ["Synthetic provenance identifiers", "Do not establish real source authenticity", "Collector signing, secure transport, and attested source identity"],
        ["Policy and verifier share configuration", "Potential common-mode specification defect", "Independent policy representation and diverse verification implementation"],
        ["Single-process runtime", "Weak fault and privilege isolation", "Process or service isolation, least privilege, sandboxing, and mTLS"],
        ["In-memory target", "Does not model vendor API or distributed-state behavior", "Test-tenant adapters, retries, idempotency, rate limits, and eventual-consistency checks"],
        ["No operator identity or approval workflow", "Cannot establish accountable human authorization", "Strong operator authentication, RBAC/ABAC, approval records, and nonrepudiation"],
    ], widths=[2.0, 2.25, 2.55], font_size=8.2)

    # 14 threat model
    heading(doc, "14.", "Threat Model and Abuse Cases", 1)
    add_table(doc, ["Threat / abuse case", "Failure sought", "POC control", "Residual risk"], [
        ["Prompt injection in ticket or log text", "Convince an agent to disable an account", "Text excluded from features; instruction detection; forced abstention", "Structured-field poisoning remains possible"],
        ["Missing or forged provenance", "Elevate untrusted evidence", "Provenance and integrity thresholds", "Synthetic IDs are not cryptographically rooted"],
        ["Correlated sensors", "Create false appearance of independent corroboration", "Source-type diversity count", "Vendor or pipeline common cause not yet modeled"],
        ["Compromised CMDB criticality", "Lower protection on mission-critical asset", "Criticality is policy input and audit field", "Needs authoritative source, change control, and independent validation"],
        ["Model overconfidence", "Cross containment threshold on weak evidence", "Evidence gate, independent verifier, criticality and action restrictions", "Threshold calibration not operationally validated"],
        ["Token theft or replay", "Execute previously approved action", "Short TTL, case/action scope, signature", "No nonce store or replay cache in v0.1"],
        ["Broker bypass", "Call target without policy", "Broker rejects missing or invalid token", "Single-process POC lacks OS-level enforcement"],
        ["False API success", "Declare containment complete when state did not change", "Post-action state verification", "Real vendor state may be delayed or partially observable"],
        ["Audit modification", "Hide unsafe decision", "Hash chain and verification", "No external timestamping or immutable storage"],
        ["Ground-truth leakage", "Artificially inflate model performance", "Separate files and runtime interfaces", "Training pipeline access still requires governance"],
    ], widths=[1.7, 1.6, 2.15, 1.35], font_size=7.9)

    # 15 implementation
    heading(doc, "15.", "Implementation Baseline", 1)
    add_para(doc, "The POC is a local Python 3.11+ package with NumPy as its only runtime dependency. The end-to-end command generates data, trains the model, runs the engine, evaluates results, verifies the audit chain, and produces JSONL, JSON, CSV, and HTML artifacts.")
    heading(doc, "15.1", "Repository Modules", 2)
    add_table(doc, ["Module", "Purpose"], [
        ["synthetic.py", "Scenario generator and separated labels"],
        ["schemas.py", "Evidence, case, ground-truth, and disposition structures"],
        ["features.py", "Allow-listed structured features and event trace"],
        ["evidence.py", "Evidence trust, conflict, and adversarial-content assessment"],
        ["model.py", "Interpretable logistic model and model artifact"],
        ["policy.py", "Deterministic four-way decision policy"],
        ["verifier.py", "Independent action-eligibility verification"],
        ["actions.py", "Authorization gate, token validation, broker, simulator, and post-action verification"],
        ["audit.py", "Append-only hash-chained audit log"],
        ["engine.py", "End-to-end decision orchestration"],
        ["metrics.py", "Evaluation metrics and CSV artifacts"],
        ["reporting.py", "Self-contained HTML dashboard"],
        ["run_poc.py", "Reproducible end-to-end entry point"],
    ], widths=[1.55, 5.25], font_size=8.8)
    heading(doc, "15.2", "Reproduction Command", 2)
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    set_cell = None
    r = p.add_run("python -m venv .venv\npip install -r requirements.txt\npython run_poc.py\nPYTHONPATH=src python -m unittest discover -s tests -v")
    set_run_font(r, name="Courier New", size=9, color=DARK)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)

    # 16 V&V
    heading(doc, "16.", "Verification and Validation", 1)
    heading(doc, "16.1", "Test Strategy", 2)
    add_table(doc, ["Category", "Focus"], [
        ["Functional", "Ingestion, assessment, policy, verification, token issuance, action, post-action verification, audit, and reporting"],
        ["Safety", "Poisoned evidence, break glass, criticality, human-only actions, missing tokens, and authorization bypass"],
        ["Security", "Token integrity, traceability, audit tamper detection, and label separation"],
        ["Model", "Discrimination, calibration, false positives, false negatives, and scenario behavior"],
        ["Failure injection", "Synthetic target-command failures and post-action detection"],
        ["Reproducibility", "Fixed seed, manifests, model artifact, policy version, and deterministic dataset generation"],
    ], widths=[1.4, 5.4], font_size=8.8)
    heading(doc, "16.2", "Automated Test Results", 2)
    add_table(doc, ["Test", "Result"], [
        ["Ground truth is absent from runtime case input", "Pass"],
        ["Small end-to-end run preserves safety invariants", "Pass"],
        ["Action broker rejects missing token", "Pass"],
        ["Audit chain detects tampering", "Pass"],
        ["Break-glass identity never auto-contains", "Pass"],
        ["Poisoned evidence forces abstention", "Pass"],
        ["Verifier rejects human-only action", "Pass"],
    ], widths=[5.35, 1.45], font_size=9)
    add_callout(doc, "Test result", "Seven of seven automated tests passed in the delivered baseline. The test suite prioritizes safety invariants over synthetic model accuracy.", fill=GREEN, border="4F8A61")

    # 17 baseline results landscape charts
    heading(doc, "17.", "Baseline Results", 1)
    add_para(doc, "The baseline used seed 20260814, 800 synthetic training cases, and 400 synthetic test cases. The model artifact, policy, decisions, audit chain, metrics, test output, and HTML dashboard are included in `outputs/baseline/`.")
    add_landscape_section(doc)
    add_figure(doc, ARCH / "05_disposition_counts.png", "Figure 6. Baseline disposition counts.", width=8.5)
    add_figure(doc, ARCH / "07_scenario_outcomes.png", "Figure 7. Decision outcome by synthetic scenario.", width=8.8)
    add_portrait_section(doc)
    heading(doc, "17.1", "Decision-Control Metrics", 2)
    add_table(doc, ["Metric", "Result"], [
        ["Expected disposition match rate", f"{100*decision['expected_disposition_match_rate']:.2f}%"],
        ["Autonomous containment count", str(decision["autonomous_containment_count"])],
        ["Autonomous containment precision", f"{100*decision['autonomous_containment_precision']:.2f}%"],
        ["False containment count", str(decision["false_containment_count"])],
        ["Compromise autonomous-containment coverage", f"{100*decision['compromise_autonomous_containment_coverage']:.2f}%"],
        ["Compromise contain-or-escalate coverage", f"{100*decision['compromise_contain_or_escalate_coverage']:.2f}%"],
        ["Investigation / abstention rate", f"{100*decision['investigation_abstention_rate']:.2f}%"],
    ], widths=[4.6, 2.2], font_size=9)
    heading(doc, "17.2", "Safety and Assurance Metrics", 2)
    add_table(doc, ["Metric", "Result"], [
        ["Unsafe automation count", str(safety["unsafe_automation_count"])],
        ["Poisoned-evidence cases", str(safety["poisoned_evidence_cases"])],
        ["Poisoned-evidence autonomous actions", str(safety["poisoned_evidence_autonomous_actions"])],
        ["Authorization without independent verifier", str(safety["authorization_without_independent_verifier"])],
        ["Evidence traceability", f"{100*safety['evidence_traceability_rate']:.2f}%"],
        ["Action-command success", f"{100*safety['action_command_success_rate']:.2f}%"],
        ["Complete post-action verification pass rate", f"{100*safety['post_action_verification_pass_rate']:.2f}%"],
        ["Audit chain", "Valid" if safety["audit_chain_valid"] else "Invalid"],
    ], widths=[4.6, 2.2], font_size=9)
    heading(doc, "17.3", "Model Metrics — Synthetic Only", 2)
    add_table(doc, ["Metric", "Result"], [
        ["Accuracy at 0.5", f"{100*model['accuracy_at_0_5']:.2f}%"],
        ["Precision at 0.5", f"{100*model['precision_at_0_5']:.2f}%"],
        ["Recall at 0.5", f"{100*model['recall_at_0_5']:.2f}%"],
        ["F1 at 0.5", f"{model['f1_at_0_5']:.4f}"],
        ["ROC AUC", f"{model['roc_auc']:.4f}"],
        ["Brier score", f"{model['brier_score']:.4f}"],
        ["10-bin expected calibration error", f"{model['expected_calibration_error_10_bin']:.4f}"],
    ], widths=[4.6, 2.2], font_size=9)
    add_callout(doc, "Interpretation", "The high synthetic model metrics are expected because the train and test sets share the same generator family. They validate the code path, not operational performance. The more consequential baseline results are zero recorded authority bypass, zero unsafe automation under the encoded generator-consistent tests, complete event-ID reference checks, and successful detection of simulated action failures.", fill=AMBER, border="C98700")

    # 18 limitations
    heading(doc, "18.", "Limitations and Open Questions", 1)
    add_table(doc, ["Gap", "Why it matters", "Resolution path"], [
        ["Synthetic-to-real distribution shift", "Real telemetry, attack behavior, and benign context will differ materially.", "Historical replay, temporal holdout, source ablation, and shadow mode"],
        ["Ground-truth uncertainty", "Incident labels are often incomplete or disputed.", "Multi-reviewer adjudication, confidence labels, and unresolved class"],
        ["Source independence", "Multiple tools may derive from the same upstream sensor or rule.", "Provenance graph and common-cause analysis"],
        ["Delayed evidence", "Important context may arrive after the initial decision.", "Event-time processing, decision revision, and evidence-aging policy"],
        ["Mission consequence model", "The same action has different costs across assets and missions.", "Asset-owner consequence workshops and action-specific loss models"],
        ["Policy common-mode error", "Policy and verifier may encode the same flawed assumption.", "Diverse verifier implementation, formal constraints, and red-team review"],
        ["Canonical-context consistency", "V0.1 trusts top-level break-glass and asset-criticality fields and does not reconcile conflicting event values.", "Versioned replay contract, cross-field validation, and fail-closed rejection"],
        ["Target-state observation", "V0.1 verifies the simulator response rather than an independent target readback.", "Read-after-write adapter contract and independently sourced observation"],
        ["Authorization and recovery hardening", "The fallback key is public; token replay control, external audit anchoring, and executable rollback are absent.", "Secrets management, one-time tokens, external audit anchor, and rollback tests"],
        ["Vendor action semantics", "Session revocation and step-up behavior vary by platform.", "Test-tenant adapters and contract tests"],
        ["Human factors", "Analyst trust, overreliance, and approval fatigue can defeat technical controls.", "Workflow study, reason codes, disagreement capture, and stop conditions"],
        ["Privacy and records handling", "Identity telemetry may contain regulated or sensitive data.", "Data minimization, pseudonymization, access control, retention, and legal review"],
    ], widths=[1.75, 2.45, 2.6], font_size=8.2)

    # 19 roadmap
    heading(doc, "19.", "Where We Are Going", 1)
    add_table(doc, ["Phase", "Data / integration", "Action level", "Exit criterion"], [
        ["0 — Concept convergence", "Problem decomposition and use-case selection", "None", "Bounded objective and product thesis"],
        ["1 — Synthetic executable baseline", "Generated multi-source evidence", "Simulator only", "Safety invariants executable and reproducible"],
        ["2 — Historical replay", "De-identified cases and canonical adapters", "None", "Evidence coverage, calibration, disagreement, and data gaps quantified"],
        ["3 — Live shadow mode", "Approved read-only live feeds", "None", "Stable schemas, analyst comparison, bounded false-containment estimate"],
        ["4 — Controlled test tenant", "Non-production identity integrations", "Change-controlled reversible actions", "Idempotency, rollback, token, and state verification validated"],
        ["5 — Limited pilot", "Small approved operational population", "Human approval initially", "Authorizing official accepts bounded residual risk"],
        ["6 — Productization", "Multi-vendor, multi-tenant, sector policy packs", "Action-class specific", "Operational support, governance, secure updates, and evidence lifecycle"],
    ], widths=[1.35, 2.35, 1.55, 1.55], font_size=8.1)
    heading(doc, "19.1", "Next 90 Days", 2)
    add_table(doc, ["Window", "Primary work", "Concrete outputs"], [
        ["Days 0–30", "Governance, data-use agreement, case taxonomy, canonical event schema, replay environment, and adjudication procedure", "Approved data contract; privacy/security plan; 50-case pilot corpus; updated threat model"],
        ["Days 31–60", "Historical ingestion, vendor adapters, evidence coverage analysis, analyst labeling, temporal replay, source ablation", "250–500 replay cases; adapter test suite; schema-gap report; initial calibration report"],
        ["Days 61–90", "Counterfactual comparison, policy tuning, independent review, shadow-mode readiness assessment", "V0.2 baseline; release-gate recommendation; shadow-mode test plan; stop conditions"],
    ], widths=[1.1, 3.2, 2.5], font_size=8.5)

    # 20 immediate sprint
    heading(doc, "20.", "Immediate Implementation Plan", 1)
    add_numbered(doc, [
        "Freeze v0.1 as a reproducible baseline and tag the source repository.",
        "Select one identity platform and one EDR source for canonical adapter discovery; do not build production credentials or write access.",
        "Obtain 50–100 de-identified cases spanning true compromise, benign administration, unresolved alerts, and known false positives.",
        "Define minimum case evidence and mark every field as required, optional, derived, sensitive, or prohibited.",
        "Create a two-reviewer adjudication process with disagreement and confidence capture.",
        "Run the v0.1 pipeline unmodified against mapped replay data to expose schema and assumption failures before tuning the model.",
        "Perform source-ablation and delayed-evidence tests to determine which sources are actually decision critical.",
        "Revise thresholds only after documenting the cost of false containment, missed compromise, investigation burden, and mission disruption.",
        "Commission an independent safety and red-team review of policy, verifier, token, and audit controls.",
        "Prepare a shadow-mode readiness decision package; maintain zero live action authority.",
    ])

    # 21 data needed
    heading(doc, "21.", "Data Required for the Next Phase", 1)
    add_para(doc, "The next phase requires data sufficient to reconstruct what was known at each decision point, not merely a final incident summary. Timestamps, source identity, and decision context are essential.")
    add_table(doc, ["Domain", "Minimum fields", "Purpose"], [
        ["Identity", "Pseudonymous subject ID; event and collection time; source IP; device ID; application; auth result; MFA method/result; session/token hash; risk flags; privilege state", "Authentication sequence, token misuse, and identity context"],
        ["Endpoint", "Device ID; event time; process and parent hashes; credential-access indicator; malware verdict; device compliance; user-session link", "Corroborate credential access and endpoint state"],
        ["Network / DNS / proxy", "Event time; source/destination; domain; protocol; VPN/proxy attribution; geo-confidence; session or device correlation", "Differentiate threat infrastructure from benign routing artifacts"],
        ["Threat intelligence", "Indicator; type; source; confidence; first/last seen; expiration; provenance", "Assess whether a match is current and independently sourced"],
        ["Asset / identity governance", "Owner; environment; mission/business function; criticality; break-glass status; privilege type; dependencies", "Determine consequence and authority boundary"],
        ["Change / maintenance", "Change ID; approved window; identities; assets; intended actions; approver; status", "Disconfirm alerts caused by authorized administration"],
        ["Travel / workforce context", "Pseudonymous subject; authorized travel window; coarse location; device assignment", "Explain geolocation anomalies while minimizing privacy exposure"],
        ["Analyst decision", "Decision time; disposition; evidence used; evidence requested; confidence; reviewer; reason code", "Measure agreement, burden, and missing evidence"],
        ["Incident ground truth", "Confirmed compromise state; discovery basis; affected identities/assets; timeline; uncertainty", "Outcome evaluation with confidence rather than false certainty"],
        ["Action and recovery", "Action; authority; request and completion time; API result; verified target state; rollback; business impact", "Validate execution, side effects, and actual outcome"],
    ], widths=[1.35, 3.65, 1.8], font_size=7.7)
    heading(doc, "21.1", "Data Governance Requirements", 2)
    add_bullets(doc, [
        "Pseudonymize users and hash session, token, device, and indicator identifiers where feasible.",
        "Retain event time and collection time separately.",
        "Preserve original source and transformation provenance.",
        "Restrict sensitive workforce context to the minimum granularity needed for the decision.",
        "Separate model-development access from evaluator labels and final adjudication.",
        "Define retention, deletion, redisclosure, incident-notification, and audit requirements before ingestion.",
    ])

    # 22 risk register
    heading(doc, "22.", "Program Risk Register", 1)
    add_table(doc, ["ID", "Risk", "Likelihood", "Impact", "Mitigation / trigger"], [
        ["R-01", "Synthetic results create false confidence", "High", "High", "Place limitation beside every result; prohibit operational claims; require replay release gate"],
        ["R-02", "Historical labels are incomplete or biased", "High", "High", "Two-reviewer adjudication, uncertainty class, evidence-based label confidence"],
        ["R-03", "Evidence cannot be reliably correlated", "Medium", "High", "Canonical identity/device/session keys; adapter quality metrics; reject low-linkage cases"],
        ["R-04", "CMDB criticality is stale", "High", "High", "Authoritative ownership, freshness, independent validation, and conservative default"],
        ["R-05", "Model and policy are tuned to the replay corpus", "Medium", "High", "Temporal holdout, blind evaluation, change control, and external review"],
        ["R-06", "Operator pressure expands action authority too early", "Medium", "Critical", "Governance charter, technical deny controls, authorizing-official gate, kill switch"],
        ["R-07", "Prompt injection migrates to structured fields", "Medium", "High", "Typed schemas, source authentication, range validation, cross-source consistency, red team"],
        ["R-08", "Vendor API action is non-idempotent or partially applied", "Medium", "High", "Test tenant, idempotency keys, retries, state verification, rollback, stop conditions"],
        ["R-09", "Identity telemetry creates privacy or labor risk", "Medium", "High", "Minimization, legal/privacy review, access controls, purpose limitation, retention"],
        ["R-10", "Common-mode defect in policy and verifier", "Medium", "High", "Diverse implementation, formal constraints, independent assessment, fault injection"],
    ], widths=[0.55, 2.05, 0.8, 0.8, 2.65], font_size=7.7)

    # 23 standards
    heading(doc, "23.", "Standards and Framework Alignment", 1)
    add_para(doc, "The POC uses standards as engineering references, not as a claim of compliance or certification. NIST AI RMF provides the Govern, Map, Measure, and Manage risk-management functions [1]. The Generative AI Profile emphasizes risks that include confabulation, information integrity, data privacy, and human-AI configuration [2]. NIST SP 800-160 provides a systems-security-engineering basis for trustworthy systems [3]. NIST SP 800-61 Rev. 3 integrates incident response into cybersecurity risk management [4]. Zero Trust Architecture reinforces explicit, resource-focused authentication and authorization without implicit trust [5]. NIST SP 800-53 provides control families relevant to access enforcement, audit, monitoring, incident handling, integrity, configuration management, and system development [6]. MITRE ATT&CK informs the valid-account, credential-access, and lateral-movement scenario logic [7–9].")
    add_table(doc, ["Reference area", "POC application"], [
        ["NIST AI RMF — Govern", "Roles, authority boundaries, policy ownership, action classification, and documented limitations"],
        ["NIST AI RMF — Map", "Use-case scope, affected actors, mission consequence, data provenance, and context"],
        ["NIST AI RMF — Measure", "Calibration, false containment, abstention, traceability, red-team cases, and post-action outcome"],
        ["NIST AI RMF — Manage", "Risk treatment, release gates, stop conditions, human escalation, and continuous review"],
        ["NIST SP 800-160", "Trustworthiness objectives, separation, assurance evidence, lifecycle engineering, and contested-environment thinking"],
        ["NIST SP 800-61 Rev. 3", "Detection, response, recovery, evidence collection, and continuous improvement"],
        ["NIST SP 800-207", "Explicit authentication and authorization at the identity, device, and resource boundary"],
        ["NIST SP 800-53 Rev. 5", "AC, AU, CA, CM, IA, IR, SA, SI, and SR control-family considerations"],
        ["MITRE ATT&CK", "Scenario coverage for valid accounts, credential access, OAuth/cloud access, and lateral movement"],
    ], widths=[2.0, 4.8], font_size=8.5)

    # 24 conclusion
    heading(doc, "24.", "Conclusion and Recommendation", 1)
    add_para(doc, "The POC establishes the beginning of a trusted AI action architecture. It demonstrates that evidence trust, model assessment, policy, verification, authorization, execution, and outcome validation can be separated and made observable. That separation is the core technical requirement for using AI in consequential cyber operations without making model confidence equivalent to authority.")
    add_para(doc, "The correct next move is not to add more autonomy. It is to add evidence realism. Historical replay should challenge the current assumptions, quantify what evidence is actually available at decision time, expose correlated sources and missing context, and determine whether the identity use case can support a defensible shadow-mode experiment.")
    add_callout(doc, "Recommendation", "Approve Phase 2 historical replay and data-contract discovery under a no-live-action constraint. Preserve the model-agnostic authority boundary. Do not authorize production actions until replay, shadow mode, test-tenant execution, independent review, and action-specific statistical release gates are complete.", fill=GREEN, border="4F8A61")

    # Appendices
    add_landscape_section(doc)
    heading(doc, "Appendix A.", "Repository and Output Inventory", 1)
    add_table(doc, ["Path", "Description"], [
        ["README.md", "Quick start, architecture summary, baseline results, and limitations"],
        ["config/policy.json", "Thresholds, evidence requirements, action classes, and safety controls"],
        ["data/*.jsonl", "Synthetic training/test cases and separate ground-truth labels"],
        ["src/adf_poc/", "POC implementation modules"],
        ["tests/", "Safety and pipeline tests"],
        ["outputs/baseline/model.json", "Trained synthetic model artifact"],
        ["outputs/baseline/decisions.jsonl", "Full decision records"],
        ["outputs/baseline/audit_chain.jsonl", "Hash-chained audit records"],
        ["outputs/baseline/metrics.json", "Full baseline metrics"],
        ["outputs/baseline/decision_summary.csv", "Case-level evaluator summary"],
        ["outputs/baseline/per_scenario_metrics.csv", "Scenario-level outcome metrics"],
        ["outputs/baseline/baseline_report.html", "Self-contained evaluation dashboard"],
        ["docs/", "Engineering plan, data card, model card, safety case, test plan, roadmap, ADRs, and diagrams"],
    ], widths=[3.1, 6.6], font_size=8.5)

    heading(doc, "Appendix B.", "Disposition and Scenario Baseline", 1)
    scenario_rows = []
    for row in METRICS["per_scenario"]:
        scenario_rows.append([
            row["scenario"], str(row["cases"]), str(row["compromised"]), str(row["no_action"]),
            str(row["investigate"]), str(row["contain_reversible"]), str(row["escalate_human"]),
            f"{100*row['expected_disposition_match_rate']:.1f}%"
        ])
    add_table(doc, ["Scenario", "Cases", "Comp.", "No action", "Investigate", "Contain", "Escalate", "Expected match"], scenario_rows, widths=[2.25, 0.65, 0.65, 0.9, 0.9, 0.8, 0.8, 1.0], font_size=7.8)

    heading(doc, "Appendix C.", "Requirements Traceability Matrix", 1)
    with (DOCS / "REQUIREMENTS_TRACEABILITY_MATRIX.csv").open(newline="", encoding="utf-8") as handle:
        reader = csv.reader(handle)
        headers = next(reader)
        rows = [row for row in reader]
    add_table(doc, headers, rows, widths=[0.85, 1.0, 3.3, 1.75, 1.65, 1.0], font_size=6.9)

    # Continue the remaining appendices in landscape. A final portrait section
    # produces a trailing blank page in some DOCX/PDF renderers.
    heading(doc, "Appendix D.", "Glossary", 1)
    add_table(doc, ["Term", "Definition"], [
        ["Abstention", "A deliberate decision not to close or act because evidence or confidence is insufficient."],
        ["Action authority", "The delegated right to cause a change in an operational system."],
        ["Action broker", "Credential-isolated component that validates authorization and calls the target interface."],
        ["Decision-grade evidence", "Evidence meeting minimum provenance, integrity, freshness, completeness, diversity, and conflict requirements for a disposition."],
        ["Evidence provenance", "Information identifying the source, collection path, and transformations associated with evidence."],
        ["Independent verifier", "Non-model component that re-evaluates action eligibility against deterministic constraints."],
        ["Model-agnostic", "Architecture that can replace the analytical model without changing the authority and safety boundary."],
        ["Poisoned evidence", "Untrusted content or structured data intended to manipulate analysis or action beyond its evidentiary role."],
        ["Post-action verification", "A check that the target reached the intended state after an action. V0.1 checks the state returned by its simulator; later phases require independent target readback."],
        ["Reversible containment", "A bounded low-impact action with defined restoration or recovery behavior; the recovery path may not recreate the original state exactly."],
        ["Shadow mode", "Live read-only operation in which recommendations are measured but no operational actions are allowed."],
    ], widths=[1.75, 5.05], font_size=8.7)

    heading(doc, "Appendix E.", "References", 1)
    references = [
        ("[1]", "National Institute of Standards and Technology, Artificial Intelligence Risk Management Framework (AI RMF 1.0), NIST AI 100-1, January 2023.", "https://doi.org/10.6028/NIST.AI.100-1"),
        ("[2]", "National Institute of Standards and Technology, Artificial Intelligence Risk Management Framework: Generative Artificial Intelligence Profile, NIST AI 600-1, July 2024.", "https://doi.org/10.6028/NIST.AI.600-1"),
        ("[3]", "National Institute of Standards and Technology, Engineering Trustworthy Secure Systems, NIST SP 800-160 Vol. 1 Rev. 1, November 2022.", "https://doi.org/10.6028/NIST.SP.800-160v1r1"),
        ("[4]", "National Institute of Standards and Technology, Incident Response Recommendations and Considerations for Cybersecurity Risk Management: A CSF 2.0 Community Profile, NIST SP 800-61 Rev. 3, April 2025.", "https://doi.org/10.6028/NIST.SP.800-61r3"),
        ("[5]", "National Institute of Standards and Technology, Zero Trust Architecture, NIST SP 800-207, August 2020.", "https://doi.org/10.6028/NIST.SP.800-207"),
        ("[6]", "National Institute of Standards and Technology, Security and Privacy Controls for Information Systems and Organizations, NIST SP 800-53 Rev. 5, including Release 5.2.0 materials.", "https://csrc.nist.gov/pubs/sp/800/53/r5/upd1/final"),
        ("[7]", "MITRE ATT&CK, Valid Accounts, Technique T1078.", "https://attack.mitre.org/techniques/T1078/"),
        ("[8]", "MITRE ATT&CK, Credential Access, Tactic TA0006.", "https://attack.mitre.org/tactics/TA0006/"),
        ("[9]", "MITRE ATT&CK, Lateral Movement, Tactic TA0008.", "https://attack.mitre.org/tactics/TA0008/"),
    ]
    for ref_id, text, url in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(ref_id + " ")
        set_run_font(r, bold=True, size=9)
        r = p.add_run(text + " ")
        set_run_font(r, size=9)
        add_hyperlink(p, url, url)

    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUT)
    print(OUT)
