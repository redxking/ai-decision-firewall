#!/usr/bin/env python3
"""Build the current engineering-status DOCX from its reviewed Markdown source.

The legacy v0.1 engineering baseline has a separate builder and remains an
immutable historical package. This builder is intentionally limited to the
current status and forward-plan document.
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import tempfile
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs" / "ENGINEERING_STATUS_AND_FORWARD_PLAN.md"
DEFAULT_OUTPUT = (
    ROOT
    / "docs"
    / "AI_Decision_Firewall_Engineering_Status_v0.3.0-alpha.1-candidate.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5F6B76"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "FFF4DD"
RISK = "FDEBEC"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
MAX_FIGURE_WIDTH = Inches(6.5)
MAX_FIGURE_HEIGHT = Inches(7.2)


def _set_run_font(
    run, *, name="Calibri", size=None, color=None, bold=None, italic=None
):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_spacing(paragraph, *, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS.items():
        tag = qn(f"w:{edge}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must total {TABLE_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            _set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_bottom_rule(paragraph, color=BLUE, size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def _add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def _add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def _add_inline(paragraph, text, *, default_bold=False):
    cursor = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            _set_run_font(run, bold=default_bold)
        token = match.group(0)
        if token.startswith("**"):
            _add_inline(paragraph, token[2:-2], default_bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(
                run,
                name="Courier New",
                size=9.5,
                color=INK,
                bold=default_bold,
            )
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            _add_hyperlink(paragraph, label, url)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _set_run_font(run, bold=default_bold)


def _create_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]

    def add_definition(num_fmt, level_text, font=None):
        abstract_id = max(abstract_ids or [0]) + 1
        abstract_ids.append(abstract_id)
        num_id = max(num_ids or [0]) + 1
        num_ids.append(num_id)

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), level_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, fmt, text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        numbering.append(num)
        return num_id

    return add_definition("bullet", "•", "Symbol"), add_definition("decimal", "%1.")


def _apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def _configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_together = True

    if "Table Citation" not in styles:
        citation = styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        citation = styles["Table Citation"]
    citation.font.name = "Calibri"
    citation.font.size = Pt(9)
    citation.font.color.rgb = RGBColor.from_string(MUTED)
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)


def _configure_page(document):
    document.settings.odd_and_even_pages_header_footer = True
    section = document.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    def configure_header(header):
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_spacing(p, before=0, after=2, line=1.0)
        left = p.add_run("AI Decision Firewall | Engineering Status")
        _set_run_font(left, size=8.5, color=MUTED, bold=True)
        _add_bottom_rule(p, color="D7DBE2", size="6")

    def configure_footer(footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_spacing(p, before=0, after=0, line=1.0)
        label = p.add_run(
            "Phase 3 simulation-only candidate - not operationally approved | Page "
        )
        _set_run_font(label, size=8, color=MUTED)
        _add_field(p, "PAGE")

    configure_header(section.header)
    configure_header(section.even_page_header)
    configure_footer(section.footer)
    configure_footer(section.even_page_footer)


def _add_page_boundary(document):
    """Start the next block on a new page without creating a blank page.

    An inline break can be pushed onto a new page when the preceding page is
    already full, which then advances again and leaves an empty page. A
    page-break-before marker remains on that naturally selected next page.
    """

    marker = document.add_paragraph()
    marker.paragraph_format.page_break_before = True
    # Do not chain the marker to a following table. LibreOffice may move that
    # pair to the next page when the prior page is full, leaving the marker on
    # an otherwise empty page. The boundary paragraph itself is sufficient.
    marker.paragraph_format.keep_with_next = False
    _set_spacing(marker, before=0, after=0, line=1.0)


def _add_masthead(document):
    spacer = document.add_paragraph()
    _set_spacing(spacer, before=0, after=10, line=1.0)
    title = document.add_paragraph()
    _set_spacing(title, before=0, after=4, line=1.0)
    run = title.add_run("ENGINEERING STATUS AND FORWARD PLAN")
    _set_run_font(run, size=23, color=INK, bold=True)
    subtitle = document.add_paragraph()
    _set_spacing(subtitle, before=0, after=14, line=1.0)
    run = subtitle.add_run(
        "AI Decision Firewall - Phase 3 simulation-only operational MVP candidate"
    )
    _set_run_font(run, size=14, color=MUTED)

    rows = [
        ("Prior evidence baseline", "0.2.0-alpha.5 | commit f7f6b5c | Phase 2.4"),
        (
            "Published predecessor",
            "0.2.0-alpha.6 | commit 854b15c | CI and Dependency Graph succeeded",
        ),
        (
            "Phase 3 candidate",
            "0.3.0-alpha.1 | 288/288 local regression | exact commit and CI pending",
        ),
        ("Review date", "2026-08-15"),
        (
            "Authority",
            "In-memory synthetic simulation only; no live connector or operational authority",
        ),
        (
            "Evidence boundary",
            "Phase 3 local CE-1 only; P2-CE-005 remains CE-0 NOT_EVALUATED",
        ),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    _set_table_geometry(table, [2160, 7200])
    for idx, (label, value) in enumerate(rows):
        _set_cell_shading(table.cell(idx, 0), LIGHT_GRAY)
        p0 = table.cell(idx, 0).paragraphs[0]
        _set_spacing(p0, before=0, after=0, line=1.0)
        _add_inline(p0, label, default_bold=True)
        p1 = table.cell(idx, 1).paragraphs[0]
        _set_spacing(p1, before=0, after=0, line=1.0)
        _add_inline(p1, value)
    rule = document.add_paragraph()
    _set_spacing(rule, before=8, after=8, line=1.0)
    _add_bottom_rule(rule, color=BLUE, size="14")


def _column_widths(column_count):
    if column_count == 1:
        return [9360]
    if column_count == 2:
        return [2520, 6840]
    if column_count == 3:
        return [2160, 3240, 3960]
    if column_count == 4:
        return [1800, 1800, 2160, 3600]
    base = TABLE_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def _add_table(document, rows):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    _set_table_geometry(table, _column_widths(len(rows[0])))
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                _set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            _set_spacing(p, before=0, after=0, line=1.0)
            _add_inline(p, value, default_bold=(r_idx == 0))
    after = document.add_paragraph()
    _set_spacing(after, before=0, after=2, line=1.0)
    return table


def _add_figure(document, source_base, source_path, alt_text, caption):
    resolved = (source_base / source_path).resolve()
    if not resolved.is_file():
        raise FileNotFoundError(resolved)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    inline = run.add_picture(str(resolved))
    natural_width = int(inline.width)
    natural_height = int(inline.height)
    scale = min(
        int(MAX_FIGURE_WIDTH) / natural_width,
        int(MAX_FIGURE_HEIGHT) / natural_height,
    )
    inline.width = int(round(natural_width * scale))
    inline.height = int(round(natural_height * scale))
    inline._inline.docPr.set("descr", alt_text)
    inline._inline.docPr.set("title", caption)
    cp = document.add_paragraph(style="Figure Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run(caption)


def _parse_table(lines, start):
    if start + 1 >= len(lines):
        return None
    first = lines[start].strip()
    divider = lines[start + 1].strip()
    if not (first.startswith("|") and re.fullmatch(r"\|?[\s:|-]+\|?", divider)):
        return None

    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        if idx != start + 1:
            cells = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
            rows.append(cells)
        idx += 1
    expected = len(rows[0])
    if any(len(row) != expected for row in rows):
        raise ValueError(f"inconsistent Markdown table near line {start + 1}")
    return rows, idx


def _render_markdown(document, source):
    bullet_id, decimal_id = _create_numbering(document)
    lines = source.read_text(encoding="utf-8").splitlines()
    paragraph_buffer = []
    page_boundary_pending = False
    in_code_block = False

    def flush_paragraph():
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer)
        p = document.add_paragraph()
        _set_spacing(p, before=0, after=6, line=1.10)
        _add_inline(p, text)
        paragraph_buffer.clear()

    index = 0
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            in_code_block = not in_code_block
            if not in_code_block:
                spacer = document.add_paragraph()
                _set_spacing(spacer, before=0, after=3, line=1.0)
            index += 1
            continue
        if in_code_block:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.right_indent = Inches(0.15)
            _set_spacing(p, before=0, after=0, line=1.0)
            run = p.add_run(raw if raw else " ")
            _set_run_font(run, name="Consolas", size=7.5, color=INK)
            index += 1
            continue
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped == "<!-- PAGE BREAK -->":
            flush_paragraph()
            next_index = index + 1
            while next_index < len(lines) and not lines[next_index].strip():
                next_index += 1
            if next_index >= len(lines) or not lines[next_index].strip().startswith(
                ("## ", "### ")
            ):
                raise ValueError("PAGE BREAK must be followed by a Markdown heading")
            page_boundary_pending = True
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("## ") or stripped.startswith("### "):
            flush_paragraph()
            level = 2 if stripped.startswith("### ") else 1
            marker_len = 4 if level == 2 else 3
            p = document.add_paragraph(style=f"Heading {level}")
            if page_boundary_pending:
                p.paragraph_format.page_break_before = True
                page_boundary_pending = False
            # Apply the heading geometry directly as well as through the style.
            # LibreOffice can collapse style-only spacing when a kept heading
            # precedes a table or a paragraph at a forced page boundary.
            _set_spacing(
                p,
                before=16 if level == 1 else 12,
                after=8 if level == 1 else 6,
                line=1.0,
            )
            next_index = index + 1
            while next_index < len(lines) and not lines[next_index].strip():
                next_index += 1
            next_is_table = next_index < len(lines) and lines[
                next_index
            ].strip().startswith("|")
            # LibreOffice needs a kept heading before ordinary prose to avoid
            # a run-in heading at a forced boundary, but keeping a heading with
            # a large table can collapse it into the table header. Select the
            # stable behavior based on the next Markdown block.
            p.paragraph_format.keep_with_next = not next_is_table
            p.paragraph_format.keep_together = True
            _add_inline(p, stripped[marker_len:])
            index += 1
            continue
        figure = re.fullmatch(r"!\[([^\]]+)\]\(([^)]+)\)", stripped)
        if figure:
            flush_paragraph()
            alt, path = figure.groups()
            _add_figure(document, source.parent, path, alt, alt)
            index += 1
            continue
        parsed_table = _parse_table(lines, index)
        if parsed_table:
            flush_paragraph()
            rows, index = parsed_table
            _add_table(document, rows)
            continue
        if stripped.startswith("> "):
            flush_paragraph()
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            _set_table_geometry(table, [TABLE_WIDTH_DXA])
            cell = table.cell(0, 0)
            fill = RISK if "Release hold" in stripped else CAUTION
            _set_cell_shading(cell, fill)
            p = cell.paragraphs[0]
            _set_spacing(p, before=0, after=0, line=1.10)
            _add_inline(p, stripped[2:])
            document.add_paragraph().paragraph_format.space_after = Pt(2)
            index += 1
            continue
        if re.match(r"^-\s+", stripped):
            flush_paragraph()
            p = document.add_paragraph()
            _set_spacing(p, before=0, after=8, line=1.167)
            _apply_numbering(p, bullet_id)
            _add_inline(p, re.sub(r"^-\s+", "", stripped))
            index += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            p = document.add_paragraph()
            _set_spacing(p, before=0, after=8, line=1.167)
            # Preserve the source number as text. LibreOffice can place the
            # next auto-number on the preceding wrapped line when a long list
            # item crosses a pagination boundary, which produced an orphaned
            # number in the paired PDF. A hanging indent keeps the intended
            # layout without relying on renderer-specific list state.
            p.paragraph_format.left_indent = Inches(0.42)
            p.paragraph_format.first_line_indent = Inches(-0.28)
            _add_inline(p, stripped)
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()
    if in_code_block:
        raise ValueError("unclosed Markdown code fence")


def _audit(document):
    section = document.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.left_margin == Inches(1)
    assert section.right_margin == Inches(1)
    assert section.top_margin == Inches(1)
    assert section.bottom_margin == Inches(1)
    assert document.styles["Normal"].font.name == "Calibri"
    assert document.styles["Normal"].font.size == Pt(11)
    assert len(document.inline_shapes) == 1
    for shape in document.inline_shapes:
        assert shape.width <= MAX_FIGURE_WIDTH
        assert shape.height <= MAX_FIGURE_HEIGHT
    for table in document.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and tbl_w.get(qn("w:w")) == str(TABLE_WIDTH_DXA)
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == str(TABLE_INDENT_DXA)


def build(source, output):
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    document.core_properties.title = (
        "AI Decision Firewall Engineering Status and Forward Plan"
    )
    document.core_properties.subject = (
        "Phase 3 simulation-only operational MVP candidate status"
    )
    document.core_properties.author = "AI Decision Firewall project"
    document.core_properties.keywords = (
        "decision assurance, simulation safety, authorization, audit"
    )
    document.core_properties.created = datetime(2026, 8, 15, tzinfo=UTC)
    document.core_properties.modified = datetime.now(UTC)
    _add_masthead(document)
    _render_markdown(document, source)
    _audit(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def convert_to_pdf(docx_path, pdf_path):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        raise RuntimeError("LibreOffice/soffice is required to build the paired PDF.")
    with tempfile.TemporaryDirectory(prefix="adf-engineering-status-") as directory:
        target = Path(directory)
        subprocess.run(
            [
                soffice,
                f"-env:UserInstallation={(target / 'profile').as_uri()}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(target),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            text=True,
            timeout=120,
        )
        rendered = target / f"{docx_path.stem}.pdf"
        if not rendered.is_file():
            raise RuntimeError("LibreOffice did not produce the expected PDF.")
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(rendered, pdf_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument(
        "--pdf-output",
        type=Path,
        help="Paired PDF path (defaults to the DOCX path with a .pdf suffix).",
    )
    args = parser.parse_args()
    output = args.output.resolve()
    pdf_output = (
        args.pdf_output.resolve() if args.pdf_output else output.with_suffix(".pdf")
    )
    build(args.source.resolve(), output)
    convert_to_pdf(output, pdf_output)
    print(output)
    print(pdf_output)


if __name__ == "__main__":
    main()
